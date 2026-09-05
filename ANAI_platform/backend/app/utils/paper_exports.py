"""Professional text/PDF/DOCX exports for generated papers."""
from __future__ import annotations
import io
from typing import Dict, Any
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem
from reportlab.lib.enums import TA_CENTER
from docx import Document


def paper_text(paper: Dict[str, Any], include_answers: bool = False) -> str:
    lines = [paper.get("exam_name", "Question Paper"), "=" * 70,
             f"Subject: {paper.get('subject', '')}", f"Duration: {paper.get('duration_minutes', '')} minutes",
             f"Total Marks: {paper.get('total_marks', '')}", ""]
    for idx, instruction in enumerate(paper.get("instructions", []) or [], 1):
        lines.append(f"{idx}. {instruction}")
    for section in paper.get("sections", []):
        lines += ["", section.get("title", "Section"), "-" * 50]
        if section.get("instructions"):
            lines.append(section["instructions"])
        for q in section.get("questions", []):
            lines.append(f"Q{q.get('question_number')}. {q.get('question_text', q.get('question', ''))}")
            for option in q.get("options", []) or []:
                lines.append(f"   {option}")
            lines.append(f"   [{q.get('marks', '')} marks]")
    if include_answers:
        lines += ["", "ANSWER KEY", "-" * 50]
        for section in paper.get("answer_key", []) or []:
            lines.append(section.get("section_title", section.get("section_id", "Section")))
            for a in section.get("answers", []):
                lines.append(f"Q{a.get('question_number')}: {a.get('answer', '')} [{a.get('marks', '')} marks]")
                if a.get("marking_scheme"): lines.append(f"Marking scheme: {a['marking_scheme']}")
    return "\n".join(lines)


def build_pdf(paper: Dict[str, Any], include_answers: bool = False) -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, rightMargin=42, leftMargin=42, topMargin=42, bottomMargin=42)
    styles = getSampleStyleSheet()
    title = styles["Title"]
    title.alignment = TA_CENTER
    body = styles["BodyText"]
    story = [Paragraph(str(paper.get("exam_name", "Question Paper")), title), Spacer(1, 12)]
    story.append(Paragraph(f"Subject: {paper.get('subject','')} &nbsp;&nbsp; Duration: {paper.get('duration_minutes','')} minutes &nbsp;&nbsp; Marks: {paper.get('total_marks','')}", body))
    story.append(Spacer(1, 12))
    for section in paper.get("sections", []):
        story.append(Paragraph(str(section.get("title", "Section")), styles["Heading2"]))
        if section.get("instructions"): story.append(Paragraph(str(section["instructions"]), body))
        for q in section.get("questions", []):
            story.append(Paragraph(f"<b>Q{q.get('question_number')}.</b> {str(q.get('question_text', q.get('question',''))).replace('&','&amp;')}", body))
            for option in q.get("options", []) or []:
                story.append(Paragraph(str(option).replace('&','&amp;'), body))
            story.append(Spacer(1, 8))
    if include_answers:
        story.append(Paragraph("Answer Key & Marking Scheme", styles["Heading1"]))
        for section in paper.get("answer_key", []) or []:
            story.append(Paragraph(str(section.get("section_title", "Section")), styles["Heading2"]))
            for a in section.get("answers", []):
                story.append(Paragraph(f"<b>Q{a.get('question_number')}</b>: {str(a.get('answer','')).replace('&','&amp;')}", body))
                if a.get("marking_scheme"): story.append(Paragraph(f"Marking: {str(a['marking_scheme']).replace('&','&amp;')}", body))
                story.append(Spacer(1, 6))
    doc.build(story)
    return buf.getvalue()


def build_docx(paper: Dict[str, Any], include_answers: bool = False) -> bytes:
    doc = Document()
    doc.add_heading(str(paper.get("exam_name", "Question Paper")), 0)
    doc.add_paragraph(f"Subject: {paper.get('subject','')} | Duration: {paper.get('duration_minutes','')} minutes | Marks: {paper.get('total_marks','')}")
    for section in paper.get("sections", []):
        doc.add_heading(str(section.get("title", "Section")), level=1)
        if section.get("instructions"): doc.add_paragraph(str(section["instructions"]))
        for q in section.get("questions", []):
            doc.add_paragraph(f"Q{q.get('question_number')}. {q.get('question_text', q.get('question',''))}")
            for option in q.get("options", []) or []: doc.add_paragraph(str(option), style="List Bullet")
    if include_answers:
        doc.add_heading("Answer Key & Marking Scheme", level=1)
        for section in paper.get("answer_key", []) or []:
            doc.add_heading(str(section.get("section_title", "Section")), level=2)
            for a in section.get("answers", []):
                doc.add_paragraph(f"Q{a.get('question_number')}: {a.get('answer','')}")
                if a.get("marking_scheme"): doc.add_paragraph(f"Marking scheme: {a['marking_scheme']}")
    buf = io.BytesIO(); doc.save(buf); return buf.getvalue()
