"""
Document parser utility for extracting text from various file formats.

This module provides functions to parse PDF, DOCX, TXT, and other document formats
for use as context in question generation.
"""

import io
import re
from backend.app.utils import get_logger

logger = get_logger(__name__)

# Try to import PDF parser
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logger.warning("PyPDF2 not installed. PDF parsing will be disabled. Install with: pip install PyPDF2")

# Try to import DOCX parser
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not installed. DOCX parsing will be disabled. Install with: pip install python-docx")


def parse_pdf_bytes(file_bytes: bytes) -> str:
    """
    Parse PDF file bytes and extract text.

    Args:
        file_bytes: PDF file as bytes

    Returns:
        str: Extracted text from PDF

    Raises:
        ImportError: If PyPDF2 is not installed
        Exception: If PDF parsing fails
    """
    if not PDF_AVAILABLE:
        raise ImportError("PyPDF2 is required for PDF parsing. Install with: pip install PyPDF2")

    try:
        pdf_file = io.BytesIO(file_bytes)
        pdf_reader = PyPDF2.PdfReader(pdf_file)

        text_parts = []
        for page_num, page in enumerate(pdf_reader.pages):
            try:
                page_text = page.extract_text()
                if page_text:
                    # Clean up the text
                    page_text = re.sub(r'\s+', ' ', page_text)  # Replace multiple spaces
                    page_text = re.sub(r'\n\s*\n', '\n\n', page_text)  # Normalize newlines
                    text_parts.append(page_text.strip())
            except Exception as e:
                logger.warning(f"Error extracting page {page_num + 1}: {str(e)}")
                continue

        full_text = "\n\n".join(text_parts)
        logger.info(f"Successfully parsed PDF: {len(full_text)} characters extracted")
        return full_text

    except Exception as e:
        logger.error(f"Error parsing PDF: {str(e)}")
        raise Exception(f"Failed to parse PDF: {str(e)}")


def parse_docx_bytes(file_bytes: bytes) -> str:
    """
    Parse DOCX file bytes and extract text.

    Args:
        file_bytes: DOCX file as bytes

    Returns:
        str: Extracted text from DOCX

    Raises:
        ImportError: If python-docx is not installed
        Exception: If DOCX parsing fails
    """
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx is required for DOCX parsing. Install with: pip install python-docx")

    try:
        docx_file = io.BytesIO(file_bytes)
        doc = Document(docx_file)

        text_parts = []

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text.strip())

        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))

        full_text = "\n\n".join(text_parts)
        logger.info(f"Successfully parsed DOCX: {len(full_text)} characters extracted")
        return full_text

    except Exception as e:
        logger.error(f"Error parsing DOCX: {str(e)}")
        raise Exception(f"Failed to parse DOCX: {str(e)}")


def parse_txt_bytes(file_bytes: bytes) -> str:
    """
    Parse text file bytes.

    Args:
        file_bytes: Text file as bytes

    Returns:
        str: Extracted text

    Raises:
        Exception: If text decoding fails
    """
    try:
        # Try different encodings
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']

        for encoding in encodings:
            try:
                text = file_bytes.decode(encoding)
                # Clean up the text
                text = re.sub(r'\r\n', '\n', text)  # Normalize line endings
                text = re.sub(r'\n\s*\n', '\n\n', text)  # Normalize multiple newlines
                logger.info(f"Successfully parsed text file with {encoding} encoding: {len(text)} characters")
                return text.strip()
            except UnicodeDecodeError:
                continue

        raise Exception("Unable to decode text file with common encodings")

    except Exception as e:
        logger.error(f"Error parsing text file: {str(e)}")
        raise Exception(f"Failed to parse text file: {str(e)}")


def parse_document_bytes(file_bytes: bytes, file_type: str) -> str:
    """
    Parse document bytes based on file type.

    Args:
        file_bytes: File as bytes
        file_type: MIME type of the file

    Returns:
        str: Extracted text from document

    Raises:
        ValueError: If file type is unsupported
        Exception: If parsing fails
    """
    logger.info(f"Parsing document of type: {file_type}")

    # Handle PDF files
    if file_type == "application/pdf":
        return parse_pdf_bytes(file_bytes)

    # Handle DOCX files
    elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return parse_docx_bytes(file_bytes)

    # Handle text files
    elif file_type in ["text/plain", "text/markdown", "text/csv"]:
        return parse_txt_bytes(file_bytes)

    # Handle unknown types
    else:
        raise ValueError(f"Unsupported file type: {file_type}")


def extract_key_sections(text: str, max_length: int = 5000) -> str:
    """
    Extract key sections from document text, prioritizing important parts.

    Args:
        text: Full document text
        max_length: Maximum length to return

    Returns:
        str: Extracted key sections
    """
    if len(text) <= max_length:
        return text

    # Try to extract introduction and conclusion
    sections = []

    # Get first 1000 chars (introduction)
    sections.append(text[:1000])

    # Get last 1000 chars (conclusion)
    sections.append(text[-1000:])

    # Look for headings and their content
    heading_pattern = r'(?m)^(#{1,3}\s+.*|.*[A-Z][A-Z\s]+:?)$'
    headings = list(re.finditer(heading_pattern, text))

    if headings:
        for heading in headings[:3]:  # Take first 3 headings
            start = heading.start()
            end = min(start + 1000, len(text))
            sections.append(text[start:end])

    # Combine unique sections
    combined = "\n\n".join(list(set(sections)))
    if len(combined) > max_length:
        combined = combined[:max_length]

    return combined


def parse_document_with_extraction(
    file_bytes: bytes,
    file_type: str,
    extract_key: bool = True,
    max_length: int = 5000
) -> dict:
    """
    Parse document and optionally extract key sections.

    Args:
        file_bytes: File as bytes
        file_type: MIME type of the file
        extract_key: Whether to extract key sections only
        max_length: Maximum length for extracted text

    Returns:
        Dict containing:
            - full_text: Complete extracted text
            - extracted_text: Key sections (if extract_key=True)
            - length: Length of extracted text
            - format: Document format
    """
    full_text = parse_document_bytes(file_bytes, file_type)

    result = {
        "full_text": full_text,
        "length": len(full_text),
        "format": file_type.split('/')[-1] if '/' in file_type else file_type
    }

    if extract_key:
        result["extracted_text"] = extract_key_sections(full_text, max_length)

    logger.info(f"Document parsed: {result['length']} characters total")
    return result